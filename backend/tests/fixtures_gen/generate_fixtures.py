"""Generates the three de-identified DOCX fixtures required by
specs/001-docx-clause-extraction/fixtures/README.md.

Run: uv run python tests/fixtures_gen/generate_fixtures.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

FIXTURES_DIR = Path(__file__).resolve().parents[3] / "specs" / "001-docx-clause-extraction" / "fixtures"


def build_normal_numbering() -> Document:
    doc = Document()
    doc.add_heading("軟體開發委外承攬合約書（範例）", level=1)
    doc.add_paragraph("立合約書人 甲方：測試股份有限公司（下稱甲方）")
    doc.add_paragraph("　　　　　乙方：範例科技有限公司（下稱乙方）")
    doc.add_paragraph("雙方同意就下列軟體開發事宜訂立本合約，條款如下：")

    doc.add_paragraph("第一條　工作範圍")
    doc.add_paragraph("乙方應依附件一所載規格，為甲方開發訂單管理系統一套。")
    doc.add_paragraph("一、系統應包含商品管理、訂單管理及報表功能。")
    doc.add_paragraph("二、如需求變更，雙方應另行協商追加報價與工期。")

    doc.add_paragraph("第二條　驗收")
    doc.add_paragraph("甲方應於乙方交付成果後十個工作天內完成驗收。")
    doc.add_paragraph("（一）逾期未回覆者，視為驗收合格。")

    doc.add_paragraph("第三條　付款方式")
    doc.add_paragraph("總價款為新臺幣一百萬元整，分三期支付。")
    doc.add_paragraph("1. 簽約後支付百分之三十。")
    doc.add_paragraph("2. 期中驗收支付百分之四十。")
    doc.add_paragraph("3. 驗收合格支付百分之三十。")

    return doc


def build_mixed_numbering() -> Document:
    doc = Document()
    doc.add_heading("軟體開發委外承攬合約書（條號混用範例）", level=1)
    doc.add_paragraph("前言：本合約係就雙方合作開發案所訂之權利義務規範，如下所述。")
    doc.add_paragraph("名詞定義：本合約所稱「成果」，係指乙方依本合約完成並交付之軟體及文件。")

    doc.add_paragraph("第壹條　智慧財產權")
    doc.add_paragraph("壹、成果之著作財產權於甲方付清全部價金後歸甲方所有。")
    doc.add_paragraph("貳、乙方保證成果未侵害第三人之智慧財產權。")

    doc.add_paragraph("第 2 條　保固")
    doc.add_paragraph("乙方應提供成果交付後一年之保固服務。")
    doc.add_paragraph("(一) 保固範圍不包含甲方自行修改所致之錯誤。")

    doc.add_paragraph("第三條　賠償責任")
    doc.add_paragraph("乙方之賠償責任上限為本合約總價款。")
    doc.add_paragraph("(1) 雙方另有約定者，從其約定。")

    return doc


def build_payment_table() -> Document:
    doc = Document()
    doc.add_heading("軟體開發委外承攬合約書（付款表格範例）", level=1)
    doc.add_paragraph("雙方同意就本開發案付款事宜訂立本合約，條款如下：")

    doc.add_paragraph("第一條　工作範圍")
    doc.add_paragraph("乙方應依附件所載規格，為甲方開發會員系統一套。")

    doc.add_paragraph("第二條　付款里程碑")
    doc.add_paragraph("雙方同意依下表所列里程碑分期付款：")

    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "期別"
    header[1].text = "里程碑"
    header[2].text = "款項比例"

    rows_data = [
        ("第一期", "簽約日", "百分之三十"),
        ("第二期", "系統驗收合格日", "百分之四十"),
        ("第三期", "上線一個月後", "百分之三十"),
    ]
    for row, (period, milestone, ratio) in zip(table.rows[1:], rows_data, strict=True):
        row.cells[0].text = period
        row.cells[1].text = milestone
        row.cells[2].text = ratio

    doc.add_paragraph("第三條　驗收")
    doc.add_paragraph("甲方應於系統上線後十個工作天內完成驗收作業。")

    return doc


def main() -> None:
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)
    build_normal_numbering().save(FIXTURES_DIR / "normal-numbering.docx")
    build_mixed_numbering().save(FIXTURES_DIR / "mixed-numbering.docx")
    build_payment_table().save(FIXTURES_DIR / "payment-table.docx")
    print(f"Fixtures written to {FIXTURES_DIR}")


if __name__ == "__main__":
    main()
